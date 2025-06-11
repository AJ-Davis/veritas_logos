"""
DOCX document parser using python-docx.
"""

from docx import Document
from docx.shared import Inches
from pathlib import Path
from typing import List

from .base_parser import BaseDocumentParser
from ..models.document import ParsedDocument, DocumentSection, DocumentFormat


class DocxParser(BaseDocumentParser):
    """Parser for DOCX files."""
    
    def __init__(self):
        super().__init__(supported_extensions={'.docx'})
    
    def get_format(self) -> DocumentFormat:
        """Return the document format this parser handles."""
        return DocumentFormat.DOCX
    
    def _parse_content(self, file_path: str) -> ParsedDocument:
        """
        Parse DOCX file content.
        
        Args:
            file_path: Path to the file to parse
            
        Returns:
            ParsedDocument instance
        """
        try:
            doc = Document(file_path)
            
            # Extract text and sections
            full_text, sections = self._extract_text_and_sections(doc)
            
            # Extract metadata
            core_props = doc.core_properties
            
            return ParsedDocument(
                content=full_text,
                sections=sections,
                raw_data={
                    "title": core_props.title,
                    "author": core_props.author,
                    "subject": core_props.subject,
                    "created": core_props.created.isoformat() if core_props.created else None,
                    "modified": core_props.modified.isoformat() if core_props.modified else None,
                    "paragraph_count": len([p for p in doc.paragraphs if p.text.strip()]),
                    "table_count": len(doc.tables)
                }
            )
            
        except Exception as e:
            return ParsedDocument(
                content="",
                errors=[f"Failed to parse DOCX file: {str(e)}"]
            )
    
    def _extract_text_and_sections(self, doc) -> tuple[str, List[DocumentSection]]:
        """
        Extract text and create sections from DOCX document.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            Tuple of (full_text, sections)
        """
        all_text = []
        sections = []
        position = 0
        
        # Process paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                section_type = self._determine_paragraph_type(para)
                
                sections.append(DocumentSection(
                    content=para.text,
                    section_type=section_type,
                    position=position,
                    metadata=self._extract_paragraph_metadata(para)
                ))
                
                all_text.append(para.text)
                position += len(para.text) + 1
        
        # Process tables
        for table_idx, table in enumerate(doc.tables):
            table_text = self._extract_table_text(table)
            if table_text.strip():
                sections.append(DocumentSection(
                    content=table_text,
                    section_type="table",
                    position=position,
                    metadata={"table_index": table_idx, "rows": len(table.rows), "cols": len(table.columns)}
                ))
                
                all_text.append(table_text)
                position += len(table_text) + 1
        
        full_text = '\n'.join(all_text)
        return full_text, sections
    
    def _determine_paragraph_type(self, paragraph) -> str:
        """
        Determine the type of a paragraph based on its style and formatting.
        
        Args:
            paragraph: python-docx Paragraph object
            
        Returns:
            Section type
        """
        # Check style name for heading styles
        if paragraph.style.name.startswith('Heading'):
            return "heading"
        
        # Check for specific style types
        style_name = paragraph.style.name.lower()
        if 'title' in style_name:
            return "title"
        elif 'subtitle' in style_name:
            return "subtitle"
        elif 'caption' in style_name:
            return "caption"
        elif 'quote' in style_name or 'block' in style_name:
            return "blockquote"
        
        # Check formatting for emphasis
        if paragraph.runs:
            # Check if entire paragraph is bold (likely a heading)
            all_bold = all(run.bold for run in paragraph.runs if run.text.strip())
            if all_bold and len(paragraph.text) < 100:
                return "heading"
        
        # Check for list-like content
        text = paragraph.text.strip()
        if text.startswith(('•', '-', '*', '◦')):
            return "list_item"
        
        # Check for numbered lists
        if text and text[0].isdigit() and ('.' in text[:10] or ')' in text[:10]):
            return "list_item"
        
        # Default to paragraph
        return "paragraph"
    
    def _extract_paragraph_metadata(self, paragraph) -> dict:
        """
        Extract metadata from a paragraph.
        
        Args:
            paragraph: python-docx Paragraph object
            
        Returns:
            Metadata dictionary
        """
        metadata = {
            "style": paragraph.style.name,
            "alignment": str(paragraph.alignment) if paragraph.alignment else None
        }
        
        # Extract formatting information
        if paragraph.runs:
            formatting = {
                "bold": any(run.bold for run in paragraph.runs if run.text.strip()),
                "italic": any(run.italic for run in paragraph.runs if run.text.strip()),
                "underline": any(run.underline for run in paragraph.runs if run.text.strip()),
            }
            metadata["formatting"] = formatting
        
        return metadata
    
    def _extract_table_text(self, table) -> str:
        """
        Extract text from a table.
        
        Args:
            table: python-docx Table object
            
        Returns:
            Table text content
        """
        table_text = []
        
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = []
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        cell_text.append(paragraph.text.strip())
                row_text.append(' '.join(cell_text))
            
            if any(cell.strip() for cell in row_text):
                table_text.append('\t'.join(row_text))
        
        return '\n'.join(table_text)
    
    def _extract_images_info(self, doc) -> List[dict]:
        """
        Extract information about images in the document.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            List of image information dictionaries
        """
        images_info = []
        
        # This is a simplified version - full image extraction would require more complex logic
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                images_info.append({
                    "target": rel.target_ref,
                    "type": rel.reltype
                })
        
        return images_info