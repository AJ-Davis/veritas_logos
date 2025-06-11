"""
PDF and DOCX Output Generators for the Veritas Logos verification system.

This module implements document generators that produce annotated PDF and DOCX files
using the annotation engine and output models. The generators apply issue highlights,
insert comments, create summaries, and add navigation aids to make verification
results actionable.

Key Features:
- Annotated PDF generation using ReportLab
- Annotated DOCX generation using python-docx
- Configurable styling based on issue severity
- Summary sections with issue overviews
- Table of contents for easy navigation
- Document metadata and generation timestamps
- Support for debate view integration
"""

import os
import uuid
import tempfile
from datetime import datetime, timezone
from typing import Dict, List, Optional, Any, Union, Tuple
from pathlib import Path
from io import BytesIO

# PDF generation
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.colors import Color, red, orange, yellow, lightblue, lightgrey
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.platypus.frames import Frame
from reportlab.platypus.doctemplate import PageTemplate, BaseDocTemplate
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY

# DOCX generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.shared import OxmlElement, qn

from .output import (
    OutputVerificationResult, DocumentAnnotation, AnnotationLayer, 
    AnnotationType, HighlightStyle, ColorScheme, OutputGenerationConfig,
    OutputFormat, DebateViewOutput
)
from .issues import UnifiedIssue, IssueSeverity
from .document import ParsedDocument


class DocumentGeneratorError(Exception):
    """Exception raised during document generation."""
    pass


class BaseDocumentGenerator:
    """
    Base class for document generators.
    
    Provides common functionality for PDF and DOCX generators including
    styling configuration, content organization, and utility methods.
    """
    
    def __init__(self, config: Optional[OutputGenerationConfig] = None):
        """Initialize the generator with configuration."""
        self.config = config or OutputGenerationConfig(output_format=OutputFormat.PDF)
        self.color_scheme = self.config.color_scheme
        self.temp_files = []  # Track temporary files for cleanup
    
    def __enter__(self):
        """Context manager entry."""
        return self
    
    def __exit__(self, exc_type, exc_val, exc_tb):
        """Context manager exit with cleanup."""
        self.cleanup()
    
    def cleanup(self):
        """Clean up temporary files."""
        for temp_file in self.temp_files:
            try:
                if os.path.exists(temp_file):
                    os.unlink(temp_file)
            except Exception:
                pass  # Ignore cleanup errors
        self.temp_files.clear()
    
    def _get_severity_color(self, severity: IssueSeverity) -> str:
        """Get color for issue severity."""
        severity_colors = {
            IssueSeverity.CRITICAL: self.color_scheme.critical,
            IssueSeverity.HIGH: self.color_scheme.high,
            IssueSeverity.MEDIUM: self.color_scheme.medium,
            IssueSeverity.LOW: self.color_scheme.low,
            IssueSeverity.INFORMATIONAL: self.color_scheme.info
        }
        return severity_colors.get(severity, self.color_scheme.info)
    
    def _should_include_issue(self, issue: UnifiedIssue) -> bool:
        """Check if issue should be included based on configuration."""
        # Define severity order (higher number = more severe)
        severity_order = {
            IssueSeverity.INFORMATIONAL: 1,
            IssueSeverity.LOW: 2,
            IssueSeverity.MEDIUM: 3,
            IssueSeverity.HIGH: 4,
            IssueSeverity.CRITICAL: 5
        }
        
        issue_level = severity_order.get(issue.severity, 1)
        min_level = severity_order.get(self.config.minimum_issue_severity, 1)
        
        return issue_level >= min_level
    
    def _truncate_text(self, text: str, max_length: int = 100) -> str:
        """Truncate text to maximum length with ellipsis."""
        if len(text) <= max_length:
            return text
        return text[:max_length-3] + "..."
    
    def _format_timestamp(self, timestamp: datetime) -> str:
        """Format timestamp for display."""
        return timestamp.strftime("%Y-%m-%d %H:%M:%S UTC")


class PdfGenerator(BaseDocumentGenerator):
    """
    PDF document generator using ReportLab.
    
    Creates annotated PDF documents with issue highlights, comments,
    summary sections, and navigation aids.
    """
    
    def __init__(self, config: Optional[OutputGenerationConfig] = None):
        """Initialize PDF generator."""
        super().__init__(config)
        self.page_width, self.page_height = letter
        self.margin = 0.75 * inch
        self.content_width = self.page_width - 2 * self.margin
        
        # Initialize styles
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
    
    def _setup_custom_styles(self):
        """Set up custom paragraph styles."""
        # Title style
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Title'],
            fontSize=24,
            spaceAfter=30,
            alignment=TA_CENTER
        ))
        
        # Issue styles by severity
        for severity in IssueSeverity:
            color_hex = self._get_severity_color(severity)
            # Convert hex to ReportLab Color
            color = Color(
                int(color_hex[1:3], 16) / 255.0,
                int(color_hex[3:5], 16) / 255.0,
                int(color_hex[5:7], 16) / 255.0
            )
            
            self.styles.add(ParagraphStyle(
                name=f'Issue{severity.value.title()}',
                parent=self.styles['Normal'],
                fontSize=10,
                leftIndent=20,
                backgroundColor=color,
                borderColor=color,
                borderWidth=1,
                borderPadding=5,
                spaceAfter=6
            ))
    
    def generate(self, output_result: OutputVerificationResult, output_path: str) -> str:
        """
        Generate annotated PDF document.
        
        Args:
            output_result: Verification results with annotations
            output_path: Path where PDF should be saved
            
        Returns:
            Path to generated PDF file
        """
        try:
            # Create document
            doc = SimpleDocTemplate(
                output_path,
                pagesize=letter,
                rightMargin=self.margin,
                leftMargin=self.margin,
                topMargin=self.margin,
                bottomMargin=self.margin
            )
            
            # Build content
            story = []
            
            # Add title page
            story.extend(self._create_title_page(output_result))
            story.append(PageBreak())
            
            # Add table of contents if enabled
            if self.config.include_table_of_contents:
                story.extend(self._create_table_of_contents(output_result))
                story.append(PageBreak())
            
            # Add executive summary if enabled
            if self.config.include_summary:
                story.extend(self._create_executive_summary(output_result))
                story.append(PageBreak())
            
            # Add main content with annotations
            story.extend(self._create_annotated_content(output_result))
            
            # Add appendices if enabled
            if self.config.include_appendices:
                story.append(PageBreak())
                story.extend(self._create_appendices(output_result))
            
            # Build PDF
            doc.build(story)
            
            return output_path
            
        except Exception as e:
            raise DocumentGeneratorError(f"Failed to generate PDF: {str(e)}")
    
    def _create_title_page(self, output_result: OutputVerificationResult) -> List:
        """Create title page content."""
        story = []
        
        # Main title
        title = Paragraph("Document Verification Report", self.styles['CustomTitle'])
        story.append(title)
        story.append(Spacer(1, 0.5 * inch))
        
        # Document information
        doc_info = [
            ["Document ID:", output_result.document_id],
            ["Verification Status:", output_result.verification_status.title()],
            ["Total Issues Found:", str(output_result.total_issues)],
            ["Overall Confidence:", f"{output_result.overall_confidence:.1%}" if output_result.overall_confidence else "N/A"],
            ["Generated At:", self._format_timestamp(output_result.generated_at)],
            ["Generator Version:", output_result.generator_version]
        ]
        
        doc_table = Table(doc_info, colWidths=[2*inch, 3*inch])
        doc_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 12),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, Color(0.7, 0.7, 0.7))
        ]))
        
        story.append(doc_table)
        story.append(Spacer(1, 0.5 * inch))
        
        # Critical issues alert
        critical_issues = output_result.critical_issues
        if critical_issues:
            alert = Paragraph(
                f"⚠️ <b>CRITICAL ALERT:</b> {len(critical_issues)} critical issue(s) found that require immediate attention.",
                self.styles['IssueCritical']
            )
            story.append(alert)
        
        return story
    
    def _create_table_of_contents(self, output_result: OutputVerificationResult) -> List:
        """Create table of contents."""
        story = []
        
        title = Paragraph("Table of Contents", self.styles['Heading1'])
        story.append(title)
        story.append(Spacer(1, 0.2 * inch))
        
        toc_items = [
            "1. Executive Summary",
            "2. Document Content with Annotations",
            "3. Issue Analysis by Severity",
            "4. Verification Details"
        ]
        
        if output_result.debate_views:
            toc_items.append("5. ACVF Debate Analysis")
        
        if self.config.include_appendices:
            toc_items.extend([
                "A. Raw Verification Data",
                "B. Methodology Notes"
            ])
        
        for item in toc_items:
            toc_para = Paragraph(item, self.styles['Normal'])
            story.append(toc_para)
            story.append(Spacer(1, 6))
        
        return story
    
    def _create_executive_summary(self, output_result: OutputVerificationResult) -> List:
        """Create executive summary section."""
        story = []
        
        title = Paragraph("Executive Summary", self.styles['Heading1'])
        story.append(title)
        story.append(Spacer(1, 0.2 * inch))
        
        # Summary statistics
        summary = output_result.generate_summary()
        
        summary_text = f"""
        This document verification report presents the results of a comprehensive analysis 
        of the submitted document. The verification process examined {summary['total_issues']} 
        potential issues across multiple verification passes.
        
        <b>Key Findings:</b>
        • Overall Verification Status: {summary['verification_status'].title()}
        • Total Issues Identified: {summary['total_issues']}
        • Overall Confidence Score: {summary['overall_confidence']:.1%}
        • Processing Time: {summary['processing_time_seconds']:.2f} seconds
        """
        
        if summary.get('acvf_debates_conducted', 0) > 0:
            summary_text += f"\n• ACVF Debates Conducted: {summary['acvf_debates_conducted']}"
        
        summary_para = Paragraph(summary_text, self.styles['Normal'])
        story.append(summary_para)
        story.append(Spacer(1, 0.3 * inch))
        
        # Issues by severity breakdown
        issues_by_severity = summary.get('issues_by_severity', {})
        if issues_by_severity:
            story.append(Paragraph("Issues by Severity:", self.styles['Heading3']))
            
            severity_data = []
            for severity, count in issues_by_severity.items():
                severity_data.append([severity.title(), str(count)])
            
            if severity_data:
                severity_table = Table(severity_data, colWidths=[2*inch, 1*inch])
                severity_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                    ('GRID', (0, 0), (-1, -1), 1, Color(0.7, 0.7, 0.7))
                ]))
                story.append(severity_table)
        
        return story
    
    def _create_annotated_content(self, output_result: OutputVerificationResult) -> List:
        """Create main document content with annotations."""
        story = []
        
        title = Paragraph("Document Content with Annotations", self.styles['Heading1'])
        story.append(title)
        story.append(Spacer(1, 0.2 * inch))
        
        # Process document content in chunks with annotations
        document = output_result.document
        content = document.content
        
        # Get all annotations sorted by position
        all_annotations = []
        for layer in output_result.annotation_layers:
            all_annotations.extend(layer.annotations)
        all_annotations.sort(key=lambda a: a.start_position)
        
        # Process content with annotations
        current_pos = 0
        
        for annotation in all_annotations:
            if not self._should_include_annotation(annotation):
                continue
            
            # Add content before annotation
            if annotation.start_position > current_pos:
                before_text = content[current_pos:annotation.start_position]
                if before_text.strip():
                    story.append(Paragraph(before_text, self.styles['Normal']))
            
            # Add annotated content
            annotated_text = content[annotation.start_position:annotation.end_position]
            story.extend(self._create_annotation_display(annotation, annotated_text))
            
            current_pos = annotation.end_position
        
        # Add remaining content
        if current_pos < len(content):
            remaining_text = content[current_pos:]
            if remaining_text.strip():
                story.append(Paragraph(remaining_text, self.styles['Normal']))
        
        return story
    
    def _should_include_annotation(self, annotation: DocumentAnnotation) -> bool:
        """Check if annotation should be included."""
        return True  # Include all annotations by default
    
    def _create_annotation_display(self, annotation: DocumentAnnotation, text: str) -> List:
        """Create display elements for an annotation."""
        story = []
        
        # Get style based on annotation
        style_name = f'Issue{annotation.style.value.title()}'
        if style_name not in self.styles:
            style_name = 'Normal'
        
        # Highlighted text
        highlighted = Paragraph(f"<b>[HIGHLIGHTED]</b> {text}", self.styles[style_name])
        story.append(highlighted)
        
        # Annotation comment
        if annotation.content:
            comment = Paragraph(f"<i>💬 {annotation.content}</i>", self.styles['Normal'])
            story.append(comment)
        
        story.append(Spacer(1, 6))
        return story
    
    def _create_appendices(self, output_result: OutputVerificationResult) -> List:
        """Create appendices section."""
        story = []
        
        title = Paragraph("Appendices", self.styles['Heading1'])
        story.append(title)
        story.append(Spacer(1, 0.2 * inch))
        
        # Appendix A: Raw verification data
        story.append(Paragraph("Appendix A: Raw Verification Data", self.styles['Heading2']))
        
        # Add detailed issue information
        for issue in output_result.issue_registry.issues:
            if self._should_include_issue(issue):
                story.extend(self._create_issue_detail(issue))
        
        return story
    
    def _create_issue_detail(self, issue: UnifiedIssue) -> List:
        """Create detailed display for an issue."""
        story = []
        
        # Issue header
        header = Paragraph(
            f"<b>{issue.issue_type.value.title()} Issue</b> (Severity: {issue.severity.value.title()})",
            self.styles['Heading3']
        )
        story.append(header)
        
        # Issue details
        details = [
            f"<b>Description:</b> {issue.description}",
            f"<b>Confidence:</b> {issue.confidence:.2f}",
            f"<b>Location:</b> Position {issue.location.start_position}-{issue.location.end_position}"
        ]
        
        if issue.remediation_suggestion:
            details.append(f"<b>Remediation:</b> {issue.remediation_suggestion}")
        
        for detail in details:
            story.append(Paragraph(detail, self.styles['Normal']))
        
        story.append(Spacer(1, 12))
        return story


class DocxGenerator(BaseDocumentGenerator):
    """
    DOCX document generator using python-docx.
    
    Creates annotated DOCX documents with issue highlights, comments,
    summary sections, and navigation aids.
    """
    
    def __init__(self, config: Optional[OutputGenerationConfig] = None):
        """Initialize DOCX generator."""
        super().__init__(config)
    
    def generate(self, output_result: OutputVerificationResult, output_path: str) -> str:
        """
        Generate annotated DOCX document.
        
        Args:
            output_result: Verification results with annotations
            output_path: Path where DOCX should be saved
            
        Returns:
            Path to generated DOCX file
        """
        try:
            # Create document
            doc = Document()
            
            # Setup document properties
            self._setup_document_properties(doc, output_result)
            
            # Setup styles
            self._setup_document_styles(doc)
            
            # Add title page
            self._add_title_page(doc, output_result)
            doc.add_page_break()
            
            # Add table of contents if enabled
            if self.config.include_table_of_contents:
                self._add_table_of_contents(doc, output_result)
                doc.add_page_break()
            
            # Add executive summary if enabled
            if self.config.include_summary:
                self._add_executive_summary(doc, output_result)
                doc.add_page_break()
            
            # Add main content with annotations
            self._add_annotated_content(doc, output_result)
            
            # Add appendices if enabled
            if self.config.include_appendices:
                doc.add_page_break()
                self._add_appendices(doc, output_result)
            
            # Save document
            doc.save(output_path)
            
            return output_path
            
        except Exception as e:
            raise DocumentGeneratorError(f"Failed to generate DOCX: {str(e)}")
    
    def _setup_document_properties(self, doc: Document, output_result: OutputVerificationResult):
        """Setup document properties and metadata."""
        core_props = doc.core_properties
        core_props.title = f"Verification Report - {output_result.document_id}"
        core_props.author = "Veritas Logos Verification System"
        core_props.subject = "Document Verification Results"
        core_props.created = output_result.generated_at
        core_props.modified = output_result.generated_at
        core_props.comments = f"Generated by Veritas Logos v{output_result.generator_version}"
    
    def _setup_document_styles(self, doc: Document):
        """Setup custom styles for the document."""
        styles = doc.styles
        
        # Create highlight styles for different severities
        for severity in IssueSeverity:
            style_name = f'Highlight{severity.value.title()}'
            
            if style_name not in styles:
                # Create character style for highlighting
                style = styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
                color_hex = self._get_severity_color(severity)
                
                # Convert hex to RGB
                r = int(color_hex[1:3], 16)
                g = int(color_hex[3:5], 16)
                b = int(color_hex[5:7], 16)
                
                # Set highlight color
                style.font.highlight_color = self._get_highlight_color_index(severity)
                style.font.color.rgb = RGBColor(0, 0, 0)  # Black text
    
    def _get_highlight_color_index(self, severity: IssueSeverity):
        """Get highlight color index for severity."""
        severity_highlights = {
            IssueSeverity.CRITICAL: WD_COLOR_INDEX.RED,
            IssueSeverity.HIGH: WD_COLOR_INDEX.YELLOW,
            IssueSeverity.MEDIUM: WD_COLOR_INDEX.TURQUOISE,
            IssueSeverity.LOW: WD_COLOR_INDEX.BRIGHT_GREEN,
            IssueSeverity.INFORMATIONAL: WD_COLOR_INDEX.GRAY_25
        }
        return severity_highlights.get(severity, WD_COLOR_INDEX.GRAY_25)
    
    def _add_title_page(self, doc: Document, output_result: OutputVerificationResult):
        """Add title page to document."""
        # Main title
        title = doc.add_heading('Document Verification Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacer
        
        # Document information table
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        
        info_data = [
            ("Document ID:", output_result.document_id),
            ("Verification Status:", output_result.verification_status.title()),
            ("Total Issues Found:", str(output_result.total_issues)),
            ("Overall Confidence:", f"{output_result.overall_confidence:.1%}" if output_result.overall_confidence else "N/A"),
            ("Generated At:", self._format_timestamp(output_result.generated_at)),
            ("Generator Version:", output_result.generator_version)
        ]
        
        for i, (label, value) in enumerate(info_data):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            
            # Make labels bold
            row.cells[0].paragraphs[0].runs[0].bold = True
        
        doc.add_paragraph()  # Spacer
        
        # Critical issues alert
        critical_issues = output_result.critical_issues
        if critical_issues:
            alert = doc.add_paragraph()
            run = alert.add_run(f"⚠️ CRITICAL ALERT: {len(critical_issues)} critical issue(s) found that require immediate attention.")
            run.bold = True
            run.font.color.rgb = RGBColor(255, 0, 0)  # Red
    
    def _add_table_of_contents(self, doc: Document, output_result: OutputVerificationResult):
        """Add table of contents."""
        doc.add_heading('Table of Contents', 1)
        
        toc_items = [
            "1. Executive Summary",
            "2. Document Content with Annotations", 
            "3. Issue Analysis by Severity",
            "4. Verification Details"
        ]
        
        if output_result.debate_views:
            toc_items.append("5. ACVF Debate Analysis")
        
        if self.config.include_appendices:
            toc_items.extend([
                "A. Raw Verification Data",
                "B. Methodology Notes"
            ])
        
        for item in toc_items:
            doc.add_paragraph(item, style='List Number')
    
    def _add_executive_summary(self, doc: Document, output_result: OutputVerificationResult):
        """Add executive summary section."""
        doc.add_heading('Executive Summary', 1)
        
        # Summary statistics
        summary = output_result.generate_summary()
        
        summary_text = f"""
        This document verification report presents the results of a comprehensive analysis 
        of the submitted document. The verification process examined {summary['total_issues']} 
        potential issues across multiple verification passes.
        """
        
        doc.add_paragraph(summary_text)
        
        # Key findings
        findings = doc.add_paragraph()
        findings.add_run("Key Findings:").bold = True
        
        findings_list = [
            f"Overall Verification Status: {summary['verification_status'].title()}",
            f"Total Issues Identified: {summary['total_issues']}",
            f"Overall Confidence Score: {summary['overall_confidence']:.1%}",
            f"Processing Time: {summary['processing_time_seconds']:.2f} seconds"
        ]
        
        if summary.get('acvf_debates_conducted', 0) > 0:
            findings_list.append(f"ACVF Debates Conducted: {summary['acvf_debates_conducted']}")
        
        for finding in findings_list:
            p = doc.add_paragraph(finding, style='List Bullet')
        
        # Issues by severity
        issues_by_severity = summary.get('issues_by_severity', {})
        if issues_by_severity:
            doc.add_heading('Issues by Severity', 3)
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Header row
            header_row = table.rows[0]
            header_row.cells[0].text = "Severity"
            header_row.cells[1].text = "Count"
            
            for cell in header_row.cells:
                cell.paragraphs[0].runs[0].bold = True
            
            # Data rows
            for severity, count in issues_by_severity.items():
                row = table.add_row()
                row.cells[0].text = severity.title()
                row.cells[1].text = str(count)
    
    def _add_annotated_content(self, doc: Document, output_result: OutputVerificationResult):
        """Add main document content with annotations."""
        doc.add_heading('Document Content with Annotations', 1)
        
        # Get all annotations sorted by position
        all_annotations = []
        for layer in output_result.annotation_layers:
            all_annotations.extend(layer.annotations)
        all_annotations.sort(key=lambda a: a.start_position)
        
        # Process content with annotations
        document = output_result.document
        content = document.content
        current_pos = 0
        
        # Create main content paragraph
        para = doc.add_paragraph()
        
        for annotation in all_annotations:
            if not self._should_include_annotation(annotation):
                continue
            
            # Add content before annotation
            if annotation.start_position > current_pos:
                before_text = content[current_pos:annotation.start_position]
                if before_text.strip():
                    para.add_run(before_text)
            
            # Add highlighted content
            annotated_text = content[annotation.start_position:annotation.end_position]
            highlighted_run = para.add_run(annotated_text)
            
            # Apply highlighting based on severity
            if annotation.related_issue_id:
                issue = self._find_issue_by_id(output_result, annotation.related_issue_id)
                if issue:
                    highlighted_run.font.highlight_color = self._get_highlight_color_index(issue.severity)
                    highlighted_run.bold = True
            
            # Add comment if present
            if annotation.content:
                comment_para = doc.add_paragraph()
                comment_run = comment_para.add_run(f"💬 {annotation.content}")
                comment_run.italic = True
                comment_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
            
            current_pos = annotation.end_position
        
        # Add remaining content
        if current_pos < len(content):
            remaining_text = content[current_pos:]
            if remaining_text.strip():
                para.add_run(remaining_text)
    
    def _should_include_annotation(self, annotation: DocumentAnnotation) -> bool:
        """Check if annotation should be included."""
        return True  # Include all annotations by default
    
    def _find_issue_by_id(self, output_result: OutputVerificationResult, issue_id: str) -> Optional[UnifiedIssue]:
        """Find issue by ID."""
        for issue in output_result.issue_registry.issues:
            if issue.issue_id == issue_id:
                return issue
        return None
    
    def _add_appendices(self, doc: Document, output_result: OutputVerificationResult):
        """Add appendices section."""
        doc.add_heading('Appendices', 1)
        
        # Appendix A: Raw verification data
        doc.add_heading('Appendix A: Raw Verification Data', 2)
        
        # Add detailed issue information
        for issue in output_result.issue_registry.issues:
            if self._should_include_issue(issue):
                self._add_issue_detail(doc, issue)
    
    def _add_issue_detail(self, doc: Document, issue: UnifiedIssue):
        """Add detailed information for an issue."""
        # Issue header
        heading = doc.add_heading(f"{issue.issue_type.value.title()} Issue (Severity: {issue.severity.value.title()})", 3)
        
        # Issue details
        details = [
            ("Description:", issue.description),
            ("Confidence:", f"{issue.confidence:.2f}"),
            ("Location:", f"Position {issue.location.start_position}-{issue.location.end_position}")
        ]
        
        if issue.remediation_suggestion:
            details.append(("Remediation:", issue.remediation_suggestion))
        
        for label, value in details:
            para = doc.add_paragraph()
            para.add_run(label).bold = True
            para.add_run(f" {value}")


# Factory function for creating generators
def create_document_generator(output_format: OutputFormat, 
                            config: Optional[OutputGenerationConfig] = None) -> BaseDocumentGenerator:
    """
    Factory function to create appropriate document generator.
    
    Args:
        output_format: The desired output format
        config: Optional configuration for the generator
        
    Returns:
        Appropriate document generator instance
        
    Raises:
        DocumentGeneratorError: If format is not supported
    """
    if output_format == OutputFormat.PDF:
        return PdfGenerator(config)
    elif output_format == OutputFormat.DOCX:
        return DocxGenerator(config)
    else:
        raise DocumentGeneratorError(f"Unsupported output format: {output_format}")


# Convenience functions
def generate_pdf_report(output_result: OutputVerificationResult, 
                       output_path: str,
                       config: Optional[OutputGenerationConfig] = None) -> str:
    """Generate PDF report from verification results."""
    generator_config = config or OutputGenerationConfig(output_format=OutputFormat.PDF)
    with PdfGenerator(generator_config) as generator:
        return generator.generate(output_result, output_path)


def generate_docx_report(output_result: OutputVerificationResult,
                        output_path: str, 
                        config: Optional[OutputGenerationConfig] = None) -> str:
    """Generate DOCX report from verification results."""
    generator_config = config or OutputGenerationConfig(output_format=OutputFormat.DOCX)
    with DocxGenerator(generator_config) as generator:
        return generator.generate(output_result, output_path)
